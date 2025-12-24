from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_bottom_border(paragraph):
    """Adds a bottom border to a paragraph (for section headers)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_jakes_resume():
    doc = Document()
    
    # 1. Setup Margins (0.5 inches narrow)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 2. Setup Default Font (Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- HEADER ---
    name_para = doc.add_paragraph()
    name_run = name_para.add_run("[Your Name]")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = "[Phone] | [Email] | [LinkedIn URL] | [GitHub URL]"
    contact_para.add_run(contact_text).font.size = Pt(10)
    
    doc.add_paragraph() # Spacer

    # --- SECTIONS ---
    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        add_bottom_border(p)
        p.paragraph_format.space_after = Pt(4)

    def add_split_line(left_text, right_text, bold_left=False, italic_left=False):
        """Adds a line with text on the left and text on the right."""
        p = doc.add_paragraph()
        # Right tab at 7.5 inches (8.5 width - 0.5 left - 0.5 right)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        
        run_left = p.add_run(left_text)
        if bold_left: run_left.bold = True
        if italic_left: run_left.italic = True
        
        p.add_run('\t') # Tab character
        
        run_right = p.add_run(right_text)
        # Dates/Locations usually plain or italic in Jake's. We'll keep plain for cleanliness.
        p.paragraph_format.space_after = Pt(0)
        return p

    # --- EDUCATION ---
    add_section_header("Education")
    
    # Uni
    add_split_line("Osmania University", "Hyderabad, India", bold_left=True)
    p = add_split_line("Bachelor of Engineering in Computer Science", "2021 – 2025", italic_left=True)
    doc.add_paragraph("CGPA: 7.8/10").paragraph_format.space_after = Pt(8)

    # High School
    add_split_line("New Middle East International School", "Riyadh, KSA", bold_left=True)
    add_split_line("High School Diploma", "2007 – 2021", italic_left=True)
    doc.add_paragraph() 

    # --- EXPERIENCE ---
    add_section_header("Experience")

    # Rootz
    add_split_line("Rootz Communication", "Aug 2024 – Aug 2025", bold_left=True)
    add_split_line("Machine Learning Intern", "Remote", italic_left=True)
    
    bullets_rootz = [
        "Developed automated Python scripts for inventory management, purchase order processing, and report generation, significantly reducing manual data entry.",
        "Visualized ad campaign metrics (impressions, clicks, CTR) by creating interactive dashboards using Agency Analytics.",
        "Managed 3+ domains on the GoDaddy platform, ensuring DNS stability and continuous uptime."
    ]
    for b in bullets_rootz:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b)
    
    doc.add_paragraph() # Spacer

    # YHills
    add_split_line("YHills", "Mar 2024 – Jun 2024", bold_left=True)
    add_split_line("Machine Learning Engineer", "Hyderabad, India", italic_left=True)
    
    bullets_yhills = [
        "Designed and deployed Python-based machine learning models for predictive analysis.",
        "Implemented Random Forest and Naïve Bayes Classifiers to optimize model accuracy for large datasets."
    ]
    for b in bullets_yhills:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b)

    doc.add_paragraph() 

    # --- PROJECTS ---
    add_section_header("Projects")

    # WordSmith
    add_split_line("WordSmith", "Python, PyTorch, Transformers, PEFT", bold_left=True)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Fine-tuned a Gemma 270m LLM to generate professional emails using HuggingFace Transformers and Parameter-Efficient Fine-Tuning (PEFT).")
    doc.add_paragraph() 

    # Automated Segregator
    add_split_line("Automated Segregator", "Python, IMAP, SMTP, Openpyxl", bold_left=True)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Engineered a Python automation script utilizing IMAP/SMTP protocols to streamline inventory management and daily report generation.")
    doc.add_paragraph() 

    # --- SKILLS ---
    add_section_header("Technical Skills")
    
    skills_p = doc.add_paragraph()
    skills_p.add_run("Languages: ").bold = True
    skills_p.add_run("Python, Java, JavaScript, TypeScript, SQL, HTML, CSS")
    
    frameworks_p = doc.add_paragraph()
    frameworks_p.add_run("Frameworks & Libraries: ").bold = True
    frameworks_p.add_run("PyTorch, TensorFlow, Scikit-learn, Pandas, NumPy, Flask, Django, React.js, Node.js")
    
    cloud_p = doc.add_paragraph()
    cloud_p.add_run("Cloud & Tools: ").bold = True
    cloud_p.add_run("Google Cloud Platform (GCP), AWS EC2, Azure, MongoDB, Firebase, Git, Linux")
    
    cert_p = doc.add_paragraph()
    cert_p.add_run("Certifications: ").bold = True
    cert_p.add_run("Power BI Bootcamp, Scientific Computing with Python (FreeCodeCamp), Google Analytics")

    # Save
    file_path = "Jakes_Resume_Saif.docx"
    doc.save(file_path)
    return file_path

# Execute
create_jakes_resume()
